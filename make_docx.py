from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER if level <= 2 else WD_ALIGN_PARAGRAPH.LEFT
    return h

def add_para(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

# ─── Title Page ───
add_heading("VISVESVARAYA TECHNOLOGICAL UNIVERSITY", 1)
add_para('"Jnana Sangama", Belagavi - 590018, Karnataka.', align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('Web3 Development Report On', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_heading('"SWIFTCLAIM — BLOCKCHAIN-BASED INSURANCE MANAGEMENT SYSTEM"', 2)
doc.add_paragraph()
add_para('Submitted by', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_heading("KUNAL RAJ S", 2)
doc.add_paragraph()
add_para('Under the Guidance of', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('[Guide Name] — Assistant Professor', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para('in partial fulfillment for the award of the degree of', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_heading("BACHELOR OF ENGINEERING", 2)
add_para('in', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_heading("COMPUTER SCIENCE AND ENGINEERING", 2)
add_para('(IoT & CYBERSECURITY INCLUDING BLOCKCHAIN)', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_heading("B.M.S. COLLEGE OF ENGINEERING", 2)
add_para('(Autonomous Institution under VTU) BENGALURU-560019', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('June 2025', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ─── Certificate ───
add_heading("CERTIFICATE", 1)
cert = doc.add_paragraph()
cert.add_run(
    'This is to certify that the project work entitled '
)
cert.add_run('"SwiftClaim — Blockchain-based Insurance Management System"').bold = True
cert.add_run(' carried out by ')
cert.add_run('Kunal Raj S').bold = True
cert.add_run(
    ' who is a bonafide student of B. M. S. College of Engineering. '
    'It is in partial fulfillment for the award of Bachelor of Engineering in '
    'Computer Science and Engineering (IoT & Cybersecurity including Blockchain) '
    'of the Visvesvaraya Technological University, Belagavi during the year 2025-26. '
    'The project report has been approved as it satisfies the academic requirements in '
    'respect of Web3 Development (23IC6AEWEB) work prescribed for the said degree.'
)

doc.add_paragraph()
add_para('Signature of the Guide', bold=True)
add_para('[Guide Name]')
add_para('Assistant Professor, Dept. of CSE(ICB), BMSCE, Bengaluru')
doc.add_paragraph()
add_para('Signature of the HOD', bold=True)
add_para('Dr Prasad G R — Professor & HOD, Dept. of CSE, BMSCE, Bengaluru')
doc.add_paragraph()
add_para('External Viva', bold=True)
add_para('Name of the Examiner: _______________________    Signature with date: _______________')

doc.add_page_break()

# ─── Declaration ───
add_heading("DECLARATION", 1)
decl = doc.add_paragraph()
decl.add_run('I, ')
decl.add_run('Kunal Raj S').bold = True
decl.add_run(
    ', student of 6th Semester, B.E, Department of Computer Science and Engineering '
    '(IoT & Cybersecurity including Blockchain), BMS College of Engineering, Bengaluru, '
    'hereby declare that, this Web3 Project entitled '
)
decl.add_run('"SwiftClaim — Blockchain-based Insurance Management System"').bold = True
decl.add_run(
    ' has been carried out by me under the guidance of [Guide Name], Assistant Professor, '
    'Department of CSE(ICB), BMS College of Engineering, Bengaluru during the academic '
    'semester February 2025 - July 2025.'
)
doc.add_paragraph()
add_para(
    'I also declare to the best of my knowledge and belief that the development reported '
    'here is not from part of any other report by any other students.'
)
doc.add_paragraph()
add_para('Signature', bold=True)
add_para('Kunal Raj S')

doc.add_page_break()

# ─── Table of Contents ───
add_heading("TABLE OF CONTENTS", 1)
toc_data = [
    ("1", "Introduction", "1"),
    ("1.1", "Overview", "1"),
    ("1.2", "Motivation", "1"),
    ("2", "Requirement Specification", "2"),
    ("2.1", "Functional Requirements", "2"),
    ("2.2", "Non-Functional Requirements", "3"),
    ("2.3", "Hardware Requirements", "3"),
    ("2.4", "Software Requirements", "3"),
    ("3", "Design", "4-5"),
    ("3.1", "Flow of Control", "4"),
    ("3.2", "Module Design", "4"),
    ("3.3", "Security Design", "5"),
    ("4", "Implementation", "6"),
    ("5", "Results", "7-9"),
    ("6", "Conclusion", "10"),
    ("7", "References", "11"),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "SI No."
hdr[1].text = "TITLE"
hdr[2].text = "Page No."
for row_data in toc_data:
    row = table.add_row().cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_paragraph()
add_heading("LIST OF FIGURES", 2)
fig_data = [
    ("5.1", "Landing Page / Connect Wallet", "7"),
    ("5.2", "User Dashboard", "7"),
    ("5.3", "Policy Marketplace", "8"),
    ("5.4", "My Policies", "8"),
    ("5.5", "Submit Claim", "9"),
    ("5.6", "Admin Dashboard", "9"),
    ("5.7", "Claims Panel (Admin)", "9"),
]
ftable = doc.add_table(rows=1, cols=3)
ftable.style = 'Table Grid'
fhdr = ftable.rows[0].cells
fhdr[0].text = "SL.no"
fhdr[1].text = "Figure"
fhdr[2].text = "Page No."
for fd in fig_data:
    fr = ftable.add_row().cells
    fr[0].text = fd[0]
    fr[1].text = fd[1]
    fr[2].text = fd[2]

doc.add_page_break()

# ─── 1. Introduction ───
add_heading("1. INTRODUCTION", 1)
doc.add_paragraph(
    'The traditional insurance industry is plagued by inefficiencies — manual claim '
    'verification, lengthy settlement cycles, opaque pricing, and a heavy reliance on '
    'intermediaries who slow down the process and increase costs. Claimants are often '
    'forced to wait weeks or months for payouts, and fraudulent claims further erode trust '
    'in the system. There is an urgent need for a transparent, automated, and trustless '
    'insurance platform.\n\n'
    'SwiftClaim addresses these challenges by leveraging Ethereum blockchain technology and '
    'smart contracts to deliver a decentralized insurance management system. The platform '
    'enables users to browse and purchase insurance policies on-chain, submit claims backed '
    'by IPFS-stored proofs, and receive instant, automated payouts — all without relying on '
    'a centralised authority.'
)

add_heading("1.1 Overview", 2)
doc.add_paragraph(
    'SwiftClaim is a full-stack decentralized application (dApp) that brings insurance '
    'policy management entirely on-chain. The platform is built on Ethereum (Sepolia testnet) '
    'and exposes two distinct role-based interfaces:\n\n'
    '• User Interface — allows policyholders to connect their MetaMask wallet, browse '
    'available policy templates on a marketplace, purchase policies by paying the premium '
    'in ETH, view their active policies, and submit insurance claims with supporting metadata '
    'stored on IPFS.\n\n'
    '• Admin Interface — enables platform administrators to create and manage policy '
    'templates, review pending claims, verify or reject them, approve payouts, and release '
    'funds directly to the claimant\'s wallet — all via smart contract calls.\n\n'
    'The frontend is built with React.js (Vite), uses Ethers.js v6 for blockchain '
    'communication, and features a modern dark-themed UI with purple branding, '
    'micro-animations, and role-based routing.'
)

add_heading("1.2 Motivation", 2)
doc.add_paragraph(
    'Insurance is one of the largest financial sectors globally, yet it continues to operate '
    'on legacy systems. Key pain points that motivated this project include:\n\n'
    '• Slow Claim Settlement: Conventional claims can take weeks due to manual verification.\n'
    '• Lack of Transparency: Policyholders have no visibility into how their premiums are '
    'used or how decisions are made.\n'
    '• Fraud Vulnerability: Centralised databases are susceptible to manipulation and '
    'fraudulent claims.\n'
    '• High Operational Cost: Intermediaries and administrative overhead add significant '
    'cost to premiums.\n\n'
    'Blockchain technology presents a compelling solution: immutable records, automated '
    'execution of agreements via smart contracts, and transparent, auditable transaction '
    'history. SwiftClaim demonstrates how a decentralized, transparent, and tamper-proof '
    'insurance platform can eliminate these pain points.'
)

doc.add_page_break()

# ─── 2. Requirements ───
add_heading("2. REQUIREMENT SPECIFICATION", 1)

add_heading("2.1 Functional Requirements", 2)
func_reqs = [
    ("Wallet Integration", "Users and admins connect MetaMask wallets to authenticate and sign transactions. Session persistence is maintained via localStorage."),
    ("Policy Marketplace", "Admins create policy templates (type, premium in ETH, coverage amount, duration, description). Users browse active templates and purchase by sending exact ETH."),
    ("Policy Management", "Users can view all their purchased policies, including policy type, coverage amount, start/expiry times, and active status."),
    ("Claim Submission", "Users submit claims against active, non-expired policies, providing claim type, incident date, description, and IPFS metadata URI. A 24-hour fraud-prevention cooldown is enforced per address."),
    ("Multi-Stage Claim Lifecycle", "Claims progress through five statuses managed by admins: PendingVerification → Verified → PayoutPending → Paid (or Rejected)."),
    ("Automated Payout", "Upon admin approval, the smart contract automatically transfers the coverage amount in ETH directly to the claimant's wallet."),
    ("Admin Controls", "Owner can grant or revoke admin roles. Admins manage policy templates and the full claim lifecycle."),
    ("Role-Based Access", "The dApp detects whether the connected address holds admin privileges and dynamically renders the appropriate interface."),
]
for title, desc in func_reqs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title + ": ").bold = True
    p.add_run(desc)

add_heading("2.2 Non-Functional Requirements", 2)
non_func = [
    ("Transparency", "All policy purchases, claim submissions, status changes, and payouts are recorded as immutable events on the blockchain."),
    ("Security", "Smart contract uses onlyOwner and onlyAdmin modifiers. Fraud detection limits claim frequency per address."),
    ("Scalability", "Mapping-based storage and event architecture designed for efficient on-chain reads. React frontend is modular and extensible."),
    ("Availability", "Core contract logic is available 24/7 as long as the Ethereum network is operational."),
    ("Usability", "Intuitive dark-themed React UI with role-based navigation, loading states, transaction feedback, and responsive design."),
]
for title, desc in non_func:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title + ": ").bold = True
    p.add_run(desc)

add_heading("2.3 Hardware Requirements", 2)
doc.add_paragraph(
    'Minimum Hardware (per developer/user):\n'
    '• Processor: Intel i5 or equivalent (AMD Ryzen 5 or above recommended)\n'
    '• RAM: 8 GB or higher\n'
    '• Storage: 256 GB SSD\n'
    '• Internet: Stable broadband connection (required for testnet transactions)'
)

add_heading("2.4 Software Requirements", 2)
doc.add_paragraph(
    'Frontend:\n'
    '• React.js 18 (via Vite 5 build tool)\n'
    '• Ethers.js v6 (blockchain interaction)\n'
    '• React Router DOM v7 (client-side routing)\n'
    '• Recharts (analytics charts)\n'
    '• Lucide React (icon library)\n\n'
    'Blockchain / Backend:\n'
    '• Solidity ^0.8.28 (Smart Contract language)\n'
    '• Hardhat (compile, test, deploy framework)\n'
    '• @nomicfoundation/hardhat-toolbox\n\n'
    'Wallet & Network:\n'
    '• MetaMask browser extension\n'
    '• Ethereum Sepolia Testnet (Chain ID: 11155111)\n\n'
    'Storage:\n'
    '• IPFS (for off-chain metadata/document URIs linked in claims)'
)

doc.add_page_break()

# ─── 3. Design ───
add_heading("3. DESIGN", 1)

add_heading("3.1 Flow of Control", 2)
doc.add_paragraph(
    'The system follows a three-tier architecture:\n\n'
    '• User Interface (UI): A React.js SPA built with Vite. Users interact through a '
    'browser with MetaMask installed. The UI detects the connected account, checks admin '
    'status by querying the smart contract, and renders the appropriate role-based view.\n\n'
    '• Smart Contract: A single Solidity contract (SwiftClaim.sol) deployed on the Ethereum '
    'Sepolia testnet. It manages all state — policy templates, user policies, claims, '
    'payouts, and admin roles.\n\n'
    '• IPFS Layer: Claim supporting documents are uploaded to IPFS by the user, and the '
    'resulting URI (metadataURI) is stored in the claim struct on-chain, ensuring '
    'tamper-proof evidence linking.\n\n'
    '• Blockchain Node Access: The frontend connects to Sepolia via MetaMask\'s injected '
    'window.ethereum provider, wrapped by Ethers.js BrowserProvider.'
)

add_heading("3.2 Module Design", 2)
modules = [
    ("Wallet & Authentication Module (Web3Context)",
     "Provides a React Context wrapping the entire application. Handles connectWallet(), "
     "account change detection, network validation, session restoration, and contract "
     "instantiation. Exposes account, isAdmin, contract, isConnecting, and networkError state."),
    ("Smart Contract Module (SwiftClaim.sol)",
     "Role Management: owner and isAdmin mapping with setAdmin() function.\n"
     "Policy Template Management: addPolicyTemplate() and updatePolicyTemplate().\n"
     "User Policy Module: buyPolicy() allows users to purchase a policy by sending ETH.\n"
     "Claims Module: submitClaim() with 24-hour fraud cooldown; status managed through "
     "verifyClaim(), approvePayout(), rejectClaim(), and releasePayout()."),
    ("User Interface Module",
     "PolicyMarketplace.jsx — Fetches and displays active policy templates; triggers buyPolicy().\n"
     "MyPolicies.jsx — Lists all policies owned by the connected address.\n"
     "SubmitClaim.jsx — Form to submit a claim tied to a specific user policy.\n"
     "UserDashboard.jsx — Overview of user stats (active policies, claims count, coverage)."),
    ("Admin Interface Module",
     "AdminDashboard.jsx — High-level contract statistics (total policies, claims, payouts).\n"
     "ClaimsPanel.jsx — Full claims management with status filter tabs and action buttons.\n"
     "PolicyEditor.jsx — Create and edit policy templates."),
]
for title, desc in modules:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(title + ":\n").bold = True
    p.add_run(desc)

add_heading("3.3 Security Design", 2)
security = [
    "Access Control Modifiers: onlyOwner restricts admin role assignment; onlyAdmin restricts policy and claim management actions.",
    "Fraud Detection: A lastClaimTime mapping enforces a CLAIM_COOLDOWN of 1 day between successive claims from any single address.",
    "Input Validation: All public functions validate input IDs, amounts (exact premium match), policy expiry, and claim status transitions on-chain.",
    "Policyholder Verification: submitClaim() verifies policy.policyholder == msg.sender to prevent unauthorized claim filing.",
    "Reentrancy Mitigation: State is updated (claim.status = Paid and policy.isActive = false) before the ETH transfer in releasePayout(), following the checks-effects-interactions pattern.",
    "Network Validation: The frontend checks the connected chain ID and alerts the user if they are not on the expected network.",
]
for item in security:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ─── 4. Implementation ───
add_heading("4. IMPLEMENTATION", 1)

impl_steps = [
    ("Development Environment Setup",
     "Install Node.js (v18+), npm, and Hardhat globally. Configure MetaMask for the "
     "Ethereum Sepolia testnet (RPC: https://rpc.sepolia.org, Chain ID: 11155111). "
     "Use VS Code with Solidity and ESLint extensions for development."),
    ("Smart Contract Development",
     "Write SwiftClaim.sol in Solidity ^0.8.28, implementing:\n"
     "• PolicyTemplate and UserPolicy structs for policy data.\n"
     "• Claim struct with ClaimStatus enum (PendingVerification, Verified, Rejected, PayoutPending, Paid).\n"
     "• Role-based access control with onlyOwner and onlyAdmin modifiers.\n"
     "• Fraud detection via lastClaimTime mapping and CLAIM_COOLDOWN constant.\n"
     "• receive() function allowing admins to fund the contract for payouts."),
    ("Deployment",
     "Configure hardhat.config.js with the Sepolia network using environment variables "
     "(SEPOLIA_RPC, PRIVATE_KEY) loaded from .env.\n"
     "Compile: npx hardhat compile\n"
     "Deploy: npx hardhat run scripts/deploy.js --network sepolia\n"
     "Copy the deployed contract address and ABI into the frontend utils/ directory."),
    ("Frontend Development",
     "Scaffold with Vite + React. Implement Web3Context using Ethers.js BrowserProvider "
     "and Contract to interface with the deployed SwiftClaim contract. Build role-based "
     "routing in App.jsx using React Router DOM. Implement all UI pages. Style using a "
     "custom CSS design system with dark background (#0d0d12), purple accents (#7c3aed), "
     "glassmorphism card effects, and smooth CSS transitions."),
    ("Testing",
     "Deploy to Sepolia and test with multiple MetaMask accounts (admin + regular user). "
     "Purchase a policy, submit a claim, then switch to admin account to verify, approve "
     "payout, and release funds. Verify on Sepolia Etherscan that all transactions are "
     "recorded with correct parameters."),
]
for title, desc in impl_steps:
    p = doc.add_paragraph()
    p.add_run("• " + title + ":\n").bold = True
    p.add_run(desc)
    doc.add_paragraph()

doc.add_page_break()

# ─── 5. Results ───
add_heading("5. RESULTS", 1)
doc.add_paragraph(
    'SwiftClaim was successfully developed and deployed as a full-stack decentralized '
    'insurance application on the Ethereum Sepolia testnet. The following key outcomes '
    'were achieved:\n\n'
    '• Smart Contract Deployment: SwiftClaim.sol was compiled and deployed to Sepolia. '
    'The contract correctly enforces all role-based access controls, policy lifecycle '
    'management, and claim status transitions.\n\n'
    '• Wallet Integration: MetaMask wallet connection functions flawlessly, with session '
    'persistence across page refreshes via localStorage.\n\n'
    '• Policy Marketplace: Admins successfully created multiple policy templates '
    '(Health, Vehicle, Property). Users purchased policies by sending the exact premium in ETH.\n\n'
    '• Claim Submission & Lifecycle: Users submitted claims against active policies. '
    'The 24-hour fraud detection cooldown correctly blocked duplicate claims. Admins '
    'moved claims through all five status stages successfully.\n\n'
    '• Automated Payouts: Upon admin triggering releasePayout(), the smart contract '
    'transferred coverage amounts in ETH directly to claimant wallets within block '
    'confirmation time (~15 seconds on Sepolia).\n\n'
    '• UI/UX: The React frontend delivered a responsive, modern dark-themed interface '
    'with real-time feedback for transaction states.\n\n'
    '• Security: All unauthorized access attempts were correctly reverted by the smart '
    'contract with appropriate error messages.'
)

doc.add_paragraph()
for fig_num, fig_title in [
    ("Fig 5.1", "Landing Page / Connect Wallet — Features the SwiftClaim logo, 'Connect Wallet to Start' button, and feature pills: Trustless, Transparent, Instant Payout, IPFS Proofs."),
    ("Fig 5.2", "User Dashboard — Shows summary cards for Active Policies, Total Claims Filed, and Total Coverage in ETH."),
    ("Fig 5.3", "Policy Marketplace — Grid of available policy template cards showing type, premium, coverage amount, and a 'Buy Policy' button."),
    ("Fig 5.4", "My Policies — Table/cards listing the user's purchased policies with policy ID, type, start date, expiry date, and active status badge."),
    ("Fig 5.5", "Submit Claim — Form with dropdowns for policy selection, claim type input, incident date picker, description textarea, and IPFS metadata URI field."),
    ("Fig 5.6", "Admin Dashboard — Statistics cards for total policy templates, user policies sold, total claims, and total ETH paid out, with a Recharts bar chart."),
    ("Fig 5.7", "Claims Panel (Admin) — Tabbed claims table filtered by status with action buttons for Verify, Approve Payout, Reject, and Release Payout."),
]:
    p = doc.add_paragraph()
    p.add_run(fig_num + " — ").bold = True
    p.add_run(fig_title)

doc.add_page_break()

# ─── 6. Conclusion ───
add_heading("6. CONCLUSION", 1)
doc.add_paragraph(
    'SwiftClaim successfully demonstrates the viability of a blockchain-based, decentralized '
    'insurance management system built on Ethereum. By replacing centralised intermediaries '
    'with a transparent, self-executing smart contract, the project addresses the core '
    'inefficiencies of traditional insurance — slow claim processing, opaque operations, '
    'and high administrative overhead.\n\n'
    'All functional requirements were met: MetaMask wallet integration, role-based access '
    'for users and admins, a policy marketplace with on-chain purchases, a multi-stage '
    'claim lifecycle, fraud detection through cooldown enforcement, and automated ETH '
    'payouts directly to claimants. The React.js frontend provides an intuitive, modern '
    'interface that lowers the barrier to entry for non-technical users.\n\n'
    'Future Enhancements:\n'
    '• Layer-2 Scaling: Deploying to Polygon, Arbitrum, or Base would reduce gas fees '
    'and improve throughput for real-world use.\n'
    '• Chainlink Oracle Integration: Enable automatic claim verification against real-world '
    'data sources (hospital APIs, police reports, weather data) for zero-touch payouts.\n'
    '• IPFS Upload Integration: Embedding a dedicated IPFS gateway (Pinata, Web3.Storage) '
    'directly in the claim submission form.\n'
    '• NFT-Based Policy Certificates: Issuing ERC-721 NFTs for each purchased policy.\n'
    '• Mobile Application: A dedicated mobile app would extend accessibility.\n'
    '• DAO Governance: Transitioning admin governance to a DAO for fully decentralized '
    'claim dispute resolution.'
)

doc.add_page_break()

# ─── 7. References ───
add_heading("7. REFERENCES", 1)
refs = [
    "Wood, G. (2014). Ethereum: A secure decentralised generalised transaction ledger. Ethereum Project Yellow Paper. https://ethereum.github.io/yellowpaper/paper.pdf",
    "Buterin, V. (2013). Ethereum White Paper: A Next-Generation Smart Contract and Decentralized Application Platform. https://ethereum.org/en/whitepaper/",
    "OpenZeppelin. (2024). Solidity Smart Contract Security Best Practices. https://docs.openzeppelin.com/contracts/",
    "Nomic Foundation. (2024). Hardhat Documentation. https://hardhat.org/docs",
    "Ethers.js Team. (2024). Ethers.js v6 Documentation. https://docs.ethers.org/v6/",
    "Ethereum Foundation. (2024). Solidity Documentation. https://docs.soliditylang.org/en/latest/",
    "Dannen, C. (2017). Introducing Ethereum and Solidity: Foundations of Cryptocurrency and Blockchain Programming for Beginners. Apress.",
    "Wood, A. (2021). Mastering Ethereum: Building Smart Contracts and DApps. O'Reilly Media.",
    "Sepolia Testnet Explorer. (2024). https://sepolia.etherscan.io/",
    "IPFS Documentation. (2024). InterPlanetary File System. https://docs.ipfs.tech/",
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f"{i}. {ref}")

# Save
out_path = r"C:\Users\Kunal Raj S\OneDrive\Desktop\WEB3\SwiftClaim_Report.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
